import io
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode

class StandaloneQRGenerator:
    def __init__(self):
        self.output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output_test")
        os.makedirs(self.output_dir, exist_ok=True)
        self.docx_path = os.path.join(self.output_dir, "test_certificate.docx")

    def create_mock_template(self):
        """Creates a dummy Word document acting as your certificate template."""
        doc = Document()
        doc.add_heading("CERTIFICATE OF COMPLETION", level=0)
        
        p1 = doc.add_paragraph("This is proudly presented to: ")
        p1.add_run("John Doe").bold = True
        
        doc.add_paragraph("\nVerify authenticity below:")
        
        # This is the exact placeholder your main script looks for
        p_qr = doc.add_paragraph("{QR_CODE}")
        p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(self.docx_path)
        print(f"1. Created mock template at: {self.docx_path}")

    def build_qr_contents(self, name, cert_id, company):
        """Step 1: Compiles the string layout that will be encoded."""
        return (
            f"Name: {name}\n"
            f"Certificate ID: {cert_id}\n"
            f"Company Name: {company}"
        )

    def generate_qr_image(self, data):
        """Step 2: Generates the QR image matrix and packages it into a memory stream."""
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_Q,
            box_size=10,
            border=4,
        )
        qr.add_data(data)
        qr.make(fit=True)
        
        # Convert to RGB image
        img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
        
        # Save image to a memory buffer instead of writing a file to the disk
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        return buffer

    def embed_qr_into_doc(self, doc_path, qr_stream, cert_id):
        """Step 3: Opens the document, tracks down the placeholder, and replaces it with the QR image."""
        doc = Document(doc_path)
        placeholder = "{QR_CODE}"
        found = False
        
        # Grab the last segment of the ID for a visual caption
        caption_text = cert_id.split('-')[-1]

        # Scan paragraphs to look for the tag
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = "" # Wipe out the text tag "{QR_CODE}"
                run = paragraph.add_run()
                qr_stream.seek(0)
                
                # Insert the binary stream directly as a picture
                run.add_picture(qr_stream, width=Inches(1.5), height=Inches(1.5))
                
                # Add the mini tracking text directly underneath it
                run.add_break()
                run.add_text(caption_text)
                found = True
                break

        if found:
            doc.save(doc_path)
            print("3. QR code successfully swapped with placeholder and saved!")
        else:
            print("❌ Placeholder tag {QR_CODE} was not found in the document.")

    def run(self):
        # 0. Setup dummy template
        self.create_mock_template()

        # Built-in Mock Values
        student_name = "John Doe"
        certificate_id = "CERT-20260616-JOHNDOE-A9B8C7"
        company_name = "HR-Automation"

        print(f"2. Processing QR data for: {student_name}...")
        
        # Run execution pipeline
        qr_data = self.build_qr_contents(student_name, certificate_id, company_name)
        qr_memory_stream = self.generate_qr_image(qr_data)
        self.embed_qr_into_doc(self.docx_path, qr_memory_stream, certificate_id)
        
        print(f"\n🎉 Success! Open the file in your directory to see it:\n{self.docx_path}")

if __name__ == "__main__":
    StandaloneQRGenerator().run()