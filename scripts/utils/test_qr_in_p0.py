import os
import qrcode
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import re

# Generate QR code
qr = qrcode.QRCode(version=1, box_size=5, border=1)
qr.add_data("https://drive.google.com/file/d/dummy_file_id/view")
qr.make(fit=True)
img = qr.make_image(fill_color="black", back_color="white")
qr_path = "scripts/output/test_qr.png"
os.makedirs("scripts/output", exist_ok=True)
img.save(qr_path)

# Load document
doc = Document("scripts/templates/Completion_Certificate.docx")

# Add run to the first paragraph (P0)
p0 = doc.paragraphs[0]
run = p0.add_run()
run.add_picture(qr_path, width=Inches(1.0), height=Inches(1.0))

# Save
docx_path = "scripts/output/test_qr_p0.docx"
doc.save(docx_path)
print("Saved DOCX to:", docx_path)

# Convert to PDF and check page count
pdf_path = "scripts/output/test_qr_p0.pdf"
try:
    convert(docx_path, pdf_path)
    print("Converted to PDF successfully!")
    with open(pdf_path, "rb") as f:
        content = f.read()
    pages = len(re.findall(b"/Type\s*/Page", content))
    print("Number of pages in PDF:", pages)
except Exception as e:
    print("PDF conversion failed:", e)
