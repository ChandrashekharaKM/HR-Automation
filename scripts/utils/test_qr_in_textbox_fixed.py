import os
import qrcode
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Load the template with the placeholder we created earlier
temp_template = "scripts/templates/Completion_Certificate_with_placeholder.docx"
test_doc = Document(temp_template)

# Generate QR code
qr = qrcode.QRCode(version=1, box_size=5, border=1)
qr.add_data("https://drive.google.com/file/d/dummy_file_id/view")
qr.make(fit=True)
img = qr.make_image(fill_color="black", back_color="white")
qr_path = "scripts/output/test_qr.png"
os.makedirs("scripts/output", exist_ok=True)
img.save(qr_path)

# Find and replace {QR_CODE} with picture ONLY inside textboxes
p_elements = test_doc.element.body.xpath("//w:p")
for p_elem in p_elements:
    # Check if the paragraph is inside a textbox
    is_textbox = len(p_elem.xpath("ancestor::w:txbxContent")) > 0
    if not is_textbox:
        continue
        
    text = "".join(node.text for node in p_elem.xpath(".//w:t"))
    if "{QR_CODE}" in text:
        from docx.text.paragraph import Paragraph
        p = Paragraph(p_elem, test_doc)
        p.text = "" # Clear text
        run = p.add_run()
        run.add_picture(qr_path, width=Inches(0.8), height=Inches(0.8))
        print("Replaced {QR_CODE} inside a textbox.")

# Save final docx
final_docx = "scripts/output/test_qr_in_textbox_fixed.docx"
test_doc.save(final_docx)
print("Saved fixed test docx:", final_docx)

# Convert to PDF
final_pdf = "scripts/output/test_qr_in_textbox_fixed.pdf"
try:
    convert(final_docx, final_pdf)
    print("Converted to PDF successfully!")
    
    # Render page 1 to PNG
    import fitz
    doc = fitz.open(final_pdf)
    page = doc.load_page(0)
    pix = page.get_pixmap()
    output_png = "scripts/output/rendered/test_qr_in_textbox_fixed.png"
    pix.save(output_png)
    print(f"Saved rendered page to {output_png}")
except Exception as e:
    print("PDF conversion failed:", e)
