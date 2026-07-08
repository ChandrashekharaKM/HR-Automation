import os
import qrcode
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

# Generate QR code
qr = qrcode.QRCode(version=1, box_size=5, border=1)
qr.add_data("https://drive.google.com/file/d/dummy_file_id/view")
qr.make(fit=True)
img = qr.make_image(fill_color="black", back_color="white")
qr_path = "scripts/output/test_qr.png"
os.makedirs("scripts/output", exist_ok=True)
img.save(qr_path)

# Load document
doc = Document("scripts/templates/Completion_Certificate.docx")

# Add a paragraph at the end
p = doc.add_paragraph()
p.alignment = 1 # Center alignment
run = p.add_run()
run.add_picture(qr_path, width=Inches(1.2), height=Inches(1.2))

# Save
docx_path = "scripts/output/test_qr_output.docx"
doc.save(docx_path)
print("Saved DOCX to:", docx_path)

# Try converting to PDF
pdf_path = "scripts/output/test_qr_output.pdf"
try:
    convert(docx_path, pdf_path)
    print("Converted to PDF successfully!")
except Exception as e:
    print("PDF conversion failed:", e)
