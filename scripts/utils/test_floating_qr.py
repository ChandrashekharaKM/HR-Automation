import os
import copy
import qrcode
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import random

# Load the original template
doc = Document("scripts/templates/Completion_Certificate.docx")
p0 = doc.paragraphs[0]

# Find the template drawing element (containing CURRENT_DATE)
drawings = p0._element.xpath(".//*[local-name()='drawing']")
template_draw = None
for draw in drawings:
    text = "".join(node.text for node in draw.xpath(".//*[local-name()='t']"))
    if "{CURRENT_DATE}" in text:
        template_draw = draw
        break

if template_draw is not None:
    print("Found template drawing!")
    new_draw = copy.deepcopy(template_draw)
    
    # 1. Assign unique IDs to prevent corruption
    doc_pr = new_draw.xpath(".//*[local-name()='docPr']")
    if doc_pr:
        new_id = str(random.randint(100000, 999999))
        doc_pr[0].set("id", new_id)
        doc_pr[0].set("name", "QR Code Text Box")
    
    # Update word anchor properties
    anchor = new_draw.xpath(".//*[local-name()='anchor']")
    if anchor:
        anchor_id = f"{random.randint(0x10000000, 0x7FFFFFFF):08X}"
        edit_id = f"{random.randint(0x10000000, 0x7FFFFFFF):08X}"
        for key in anchor[0].attrib.keys():
            if "anchorId" in key:
                anchor[0].set(key, anchor_id)
            elif "editId" in key:
                anchor[0].set(key, edit_id)
                
    # 2. Update positioning (EMUs)
    # H = 2.4 in (with 1.5 in margin, this places it at 3.9 in from left edge)
    # V = 5.25 in (with 0.96 in margin, this is 6.2 in from top edge)
    # Size = 1.0 in
    h_offset = int(2.4 * 914400)
    v_offset = int(5.25 * 914400)
    extent_val = int(1.1 * 914400)
    
    h_pos = new_draw.xpath(".//*[local-name()='positionH']/*[local-name()='posOffset']")
    if h_pos:
        h_pos[0].text = str(h_offset)
        
    v_pos = new_draw.xpath(".//*[local-name()='positionV']/*[local-name()='posOffset']")
    if v_pos:
        v_pos[0].text = str(v_offset)
        
    extent = new_draw.xpath(".//*[local-name()='extent']")
    if extent:
        extent[0].set("cx", str(extent_val))
        extent[0].set("cy", str(extent_val))
        
    a_ext = new_draw.xpath(".//*[local-name()='ext']")
    for ext in a_ext:
        ext.set("cx", str(extent_val))
        ext.set("cy", str(extent_val))
        
    # 3. Set text to {QR_CODE}
    t_elements = new_draw.xpath(".//*[local-name()='t']")
    if t_elements:
        t_elements[0].text = "{QR_CODE}"
        for t_el in t_elements[1:]:
            t_el.text = ""
            
    # Append new drawing run
    run = p0.add_run()
    run._r.append(new_draw)
    print("Injected new floating drawing element with unique IDs.")
else:
    print("Template drawing not found.")

# Save temporary docx
temp_docx = "scripts/output/test_floating_qr_temp.docx"
doc.save(temp_docx)

# Load the temporary docx to perform the placeholder replacement
doc_replaced = Document(temp_docx)

# Generate QR code
qr = qrcode.QRCode(version=1, box_size=5, border=1)
qr.add_data("https://drive.google.com/file/d/dummy_file_id/view")
qr.make(fit=True)
img = qr.make_image(fill_color="black", back_color="white")
qr_path = "scripts/output/test_qr.png"
img.save(qr_path)

# Find and replace {QR_CODE} with picture ONLY inside textboxes
p_elements = doc_replaced.element.body.xpath(".//*[local-name()='p']")
for p_elem in p_elements:
    is_textbox = len(p_elem.xpath("ancestor::*[local-name()='txbxContent']")) > 0
    if not is_textbox:
        continue
    text = "".join(node.text for node in p_elem.xpath(".//*[local-name()='t']"))
    if "{QR_CODE}" in text:
        from docx.text.paragraph import Paragraph
        p = Paragraph(p_elem, doc_replaced)
        p.text = ""
        run = p.add_run()
        run.add_picture(qr_path, width=Inches(0.9), height=Inches(0.9))
        print("Replaced {QR_CODE} in textbox with picture.")

# Save final
final_docx = "scripts/output/test_floating_qr.docx"
doc_replaced.save(final_docx)

# Convert and Render
final_pdf = "scripts/output/test_floating_qr.pdf"
try:
    convert(final_docx, final_pdf)
    print("Converted to PDF successfully!")
    
    import fitz
    f_doc = fitz.open(final_pdf)
    page = f_doc.load_page(0)
    pix = page.get_pixmap()
    output_png = "scripts/output/rendered/test_floating_qr.png"
    pix.save(output_png)
    print(f"Saved rendered page to {output_png}")
except Exception as e:
    print("PDF conversion failed:", e)
