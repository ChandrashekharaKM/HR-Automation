from docx import Document

doc = Document("scripts/templates/Completion_Certificate.docx")
print("Number of paragraphs:", len(doc.paragraphs))
print("Number of tables:", len(doc.tables))
print("Number of inline shapes:", len(doc.inline_shapes))

# Let's inspect all elements recursively
def print_elements(element, depth=0):
    indent = "  " * depth
    for child in element:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        text = child.text.strip() if child.text else ""
        # print some info about key elements
        if tag in ["p", "tbl", "tr", "tc", "txbxContent", "drawing", "t"]:
            print(f"{indent}<{tag}> {text[:30]}")
            print_elements(child, depth + 1)
        else:
            print_elements(child, depth)

print("\n--- Recursive Tree Structure ---")
print_elements(doc.element.body)
