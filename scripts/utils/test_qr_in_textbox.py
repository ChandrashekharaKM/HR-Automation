import os
import qrcode
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 1. Modify the template to add {QR_CODE} placeholder in the description textbox
doc = Document("scripts/templates/Completion_Certificate.docx")
p_elements = doc.element.body.xpath("//w:p")

# Let's find the paragraph inside the textbox that contains "continued success."
target_p_elem = None
for p_elem in p_elements:
    text = "".join(node.text for node in p_elem.xpath(".//w:t"))
    if "continued success." in text and len(p_elem.xpath("ancestor::w:txbxContent")) > 0:
        target_p_elem = p_elem
        break

if target_p_elem is not None:
    print("Found target paragraph inside textbox!")
    # Let's create a new paragraph inside the same parent (txbxContent)
    parent = target_p_elem.getparent()
    
    # We will create a new w:p element
    new_p = OxmlElement('w:p')
    # Set paragraph alignment to center (w:jc w:val="center")
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)
    
    # Add a run with {QR_CODE}
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = "{QR_CODE}"
    new_r.append(new_t)
    new_p.append(new_r)
    
    # Append after target_p_elem
    target_idx = parent.index(target_p_elem)
    parent.insert(target_idx + 1, new_p)
    print("Inserted {QR_CODE} placeholder paragraph after target paragraph.")
else:
    print("Could not find target paragraph inside textbox.")

# Save modified template
temp_template = "scripts/templates/Completion_Certificate_with_placeholder.docx"
doc.save(temp_template)
print("Saved temporary template with placeholder:", temp_template)

# 2. Let's do the generation test
# Load the template we just created
test_doc = Document(temp_template)

# Generate QR code
qr = qrcode.QRCode(version=1, box_size=5, border=1)
qr.add_data("https://drive.google.com/file/d/dummy_file_id/view")
qr.make(fit=True)
img = qr.make_image(fill_color="black", back_color="white")
qr_path = "scripts/output/test_qr.png"
os.makedirs("scripts/output", exist_ok=True)
img.save(qr_path)

# Find and replace {QR_CODE} with picture
p_elements = test_doc.element.body.xpath("//w:p")
qr_replaced = False
for p_elem in p_elements:
    text = "".join(node.text for node in p_elem.xpath(".//w:t"))
    if "{QR_CODE}" in text:
        # Wrap in python-docx Paragraph
        from docx.text.paragraph import Paragraph
        p = Paragraph(p_elem, test_doc)
        # Clear existing text
        p.text = ""
        # Add picture
        run = p.add_run()
        run.add_picture(qr_path, width=Inches(0.8), height=Inches(0.8))
        qr_replaced = True
        print("Replaced {QR_CODE} with picture.")

# Save final docx
final_docx = "scripts/output/test_qr_in_textbox.docx"
test_doc.save(final_docx)
print("Saved test docx:", final_docx)

# Convert to PDF and check page count
final_pdf = "scripts/output/test_qr_in_textbox.pdf"
try:
    convert(final_docx, final_pdf)
    print("Converted to PDF successfully!")
    with open(final_pdf, "rb") as f:
        content = f.read()
    pages = len(re.findall(b"/Type\s*/Page", content))
    print("Number of pages in PDF:", pages)
except Exception as e:
    print("PDF conversion failed:", e)
