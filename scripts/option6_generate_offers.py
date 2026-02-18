import os
import re
import gspread
import pickle
from docx import Document
from docx2pdf import convert
from datetime import datetime
from dotenv import load_dotenv
from oauth2client.service_account import ServiceAccountCredentials

# Google Drive API Imports
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request

G, R, Y, B, C, W = '\033[92m', '\033[91m', '\033[93m', '\033[94m', '\033[96m', '\033[0m'

class OfferLetterGenerator:
    def __init__(self):
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        load_dotenv(os.path.join(self.script_dir, "..", ".env"))
        
        self.creds_file = os.path.join(self.script_dir, "service_account.json")
        self.client_secret_file = os.path.join(self.script_dir, "client_secret.json") 
        self.token_file = os.path.join(self.script_dir, "token.json") 
        
        self.sheet_url = os.getenv("REGISTRATION_SHEET_URL")
        self.drive_folder_id = os.getenv("OFFER_LETTERS_DRIVE_FOLDER_ID")

        self.template_docx = os.path.join(self.script_dir, "templates", "offer_template.docx")
        self.output_dir = os.path.join(self.script_dir, "output", "offer_letters")
        os.makedirs(self.output_dir, exist_ok=True)
        
        self.drive_service = None 
        self.sheet_instance = None

    def get_drive_service(self):
        creds = None
        if os.path.exists(self.token_file):
            with open(self.token_file, 'rb') as token:
                creds = pickle.load(token)
        
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    creds.refresh(Request())
                except:
                    if os.path.exists(self.token_file): os.remove(self.token_file)
                    return self.get_drive_service()
            else:
                if not os.path.exists(self.client_secret_file):
                    print(f"{R}❌ Error: 'client_secret.json' missing.{W}")
                    return None
                
                print(f"{Y}🔓 Opening browser for Google Login...{W}")
                flow = InstalledAppFlow.from_client_secrets_file(
                    self.client_secret_file, ['https://www.googleapis.com/auth/drive.file']
                )
                creds = flow.run_local_server(port=0)
            
            with open(self.token_file, 'wb') as token:
                pickle.dump(creds, token)

        return build('drive', 'v3', credentials=creds)

    def connect_services(self):
        try:
            scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
            creds = ServiceAccountCredentials.from_json_keyfile_name(self.creds_file, scope)
            client = gspread.authorize(creds)
            self.drive_service = self.get_drive_service()

            if not self.sheet_url:
                print(f"{R}❌ Error: REGISTRATION_SHEET_URL not found in .env{W}")
                return None
            
            spreadsheet = client.open_by_url(self.sheet_url)
            gid_match = re.search(r'gid=([0-9]+)', self.sheet_url)
            if gid_match:
                target_gid = int(gid_match.group(1))
                for sheet in spreadsheet.worksheets():
                    if sheet.id == target_gid:
                        self.sheet_instance = sheet
                        return sheet

            self.sheet_instance = spreadsheet.get_worksheet(0)
            return self.sheet_instance
        except Exception as e:
            print(f"{R}❌ Connection Error: {e}{W}")
            return None

    def upload_to_drive(self, file_path, file_name):
        if not self.drive_service or not self.drive_folder_id:
            print(f"{Y}⚠️  Drive Service not connected or Folder ID missing.{W}")
            return

        try:
            print(f"{Y}☁️  Uploading to Google Drive...{W}", end=" ")
            file_metadata = {'name': file_name, 'parents': [self.drive_folder_id]}
            media = MediaFileUpload(file_path, mimetype='application/pdf')
            file = self.drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
            print(f"{G}Done! (File ID: {file.get('id')}){W}")
            return file.get('id')
        except Exception as e:
            print(f"\n{R}❌ Drive Upload Failed: {e}{W}")

    def update_sheet_start_date(self, row_idx, new_date):
        try:
            headers = self.sheet_instance.row_values(1)
            col_idx = next((i for i, h in enumerate(headers, 1) if "start" in h.lower() and "date" in h.lower()), None)
            if col_idx:
                self.sheet_instance.update_cell(row_idx, col_idx, new_date)
            else:
                print(f"{Y}⚠️  Could not find 'Start_Date' column to update.{W}")
        except Exception as e:
            print(f"{R}⚠️  Sheet update failed: {e}{W}")

    def replace_placeholders(self, doc, data):
        for paragraph in doc.paragraphs:
            if self._contains_placeholder(paragraph.text, data):
                self._replace_in_paragraph(paragraph, data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if self._contains_placeholder(paragraph.text, data):
                            self._replace_in_paragraph(paragraph, data)

    def _contains_placeholder(self, text, data):
        return any(f"{{{key}}}" in text for key in data.keys())

    def _replace_in_paragraph(self, paragraph, data):
        original_text = paragraph.text
        new_text = original_text
        for key, value in data.items():
            placeholder = f"{{{key}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
        if new_text == original_text: return
        
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_props = {'name': first_run.font.name, 'size': first_run.font.size, 'bold': first_run.font.bold, 'italic': first_run.font.italic, 'underline': first_run.font.underline, 'color': first_run.font.color.rgb if first_run.font.color else None}
            for _ in range(len(paragraph.runs)): paragraph._element.remove(paragraph.runs[0]._element)
            new_run = paragraph.add_run(new_text)
            new_run.font.name = font_props['name']; new_run.font.size = font_props['size']; new_run.font.bold = font_props['bold']; new_run.font.italic = font_props['italic']; new_run.font.underline = font_props['underline']
            if font_props['color']: new_run.font.color.rgb = font_props['color']

    def _smart_get_value(self, candidate, keywords):
        for key, value in candidate.items():
            for kw in keywords:
                if kw.lower() in key.lower():
                    if value and str(value).strip():
                        return str(value).strip()
        return None

    def generate_offer_pdf(self, candidate, manual_data, row_idx):
        full_name = self._smart_get_value(candidate, ["name", "full name", "student name", "candidate"])
        if not full_name: full_name = "Candidate"
        
        candidate_id = self._smart_get_value(candidate, ["id", "roll", "reg"])
        
        data = {
            "FULL_NAME": full_name, "ROLE": manual_data['role'],
            "DATE": manual_data['date'], "PLACE": manual_data['place'],
            "START_DATE": manual_data['start_date']
        }
        
        doc = Document(self.template_docx)
        self.replace_placeholders(doc, data)
        
        safe_name = "".join([c if c.isalnum() or c == '_' else "_" for c in full_name])
        
        if candidate_id:
            safe_id = "".join([c if c.isalnum() or c == '_' else "" for c in candidate_id])
            pdf_filename = f"Offer_{safe_name}_{safe_id}.pdf"
        else:
            pdf_filename = f"Offer_{safe_name}.pdf"

        docx_path = os.path.join(self.output_dir, f"Offer_{safe_name}.docx")
        pdf_path = os.path.join(self.output_dir, pdf_filename)
        
        doc.save(docx_path)
        try:
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                os.remove(docx_path)
                return pdf_path, pdf_filename
            return docx_path, None
        except Exception as e:
            print(f"{R}❌ PDF Error: {e}{W}")
            return docx_path, None

    def _smart_get_start_date(self, candidate):
        keywords = ["start_date", "start date", "joining_date", "doj", "expected start"]
        val = self._smart_get_value(candidate, keywords)
        return val if val else "TBD"

    def run_process(self):
        sheet = self.connect_services()
        if not sheet: return
        
        candidates = sheet.get_all_records()
        hired_count = sum(1 for c in candidates if str(c.get("Status", "")).strip() == "Hired")

        if hired_count == 0:
            print(f"{Y}⚠️  No candidates marked 'Hired' found.{W}")
            return

        print(f"{Y}🔍 Found {hired_count} Hired Candidate(s)...{W}")

        for idx, candidate in enumerate(candidates):
            status = str(candidate.get("Status", "")).strip()
            if status != "Hired": continue 
            
            name = self._smart_get_value(candidate, ["name", "full name", "candidate"]) or "Unknown"
            print(f"\n{B}[{idx+1}/{len(candidates)}] {name} ({G}{status}{W}){W}")
            
            choice = input(f"Generate? (y/n/exit): ").strip().lower()
            if choice == "exit": break
            if choice != "y": continue
            
            today = datetime.now().strftime("%B %d, %Y")
            m_date = input(f"📅 Date [{today}]: ").strip() or today
            m_place = input(f"📍 Place [Bengaluru]: ").strip() or "Bengaluru"
            m_role = input(f"💼 Role [Software Developer - Intern]: ").strip() or "Software Developer - Intern"
            
            # --- UPDATED: Ask User Explicitly ---
            fetched_start = self._smart_get_start_date(candidate)
            m_start = "TBD"

            if fetched_start and fetched_start != "TBD":
                print(f"   🗓️  Found Start Date in Sheet : {C}{fetched_start}{W}")
                use_sheet = input(f"   👉 Use this date? (y/n): ").strip().lower()
                
                if use_sheet == 'y':
                    m_start = fetched_start
                else:
                    m_start = input(f"   ✍️  Enter Start Date: ").strip()
            else:
                m_start = input(f"   ✍️  Enter Start Date: ").strip()

            # Update sheet if user typed a new date
            if m_start and m_start != fetched_start:
                self.update_sheet_start_date(idx + 2, m_start)
            # ------------------------------------
            
            manual_inputs = {'date': m_date, 'place': m_place, 'role': m_role, 'start_date': m_start}

            try:
                print(f"{Y}⚙️  Generating...{W}", end=" ", flush=True)
                output_path, file_name = self.generate_offer_pdf(candidate, manual_inputs, idx+2)
                print(f"{G}✅ {output_path}{W}")
                
                if output_path and file_name and output_path.endswith('.pdf'):
                    self.upload_to_drive(output_path, file_name)

            except Exception as e:
                print(f"{R}❌ Failed: {e}{W}")

def main():
    print(f"{B}📄 SwipeGen Offer Letter Generator{W}\n")
    generator = OfferLetterGenerator()
    generator.run_process()
    print(f"\n{G}✨ Done!{W}")

if __name__ == "__main__":
    main()